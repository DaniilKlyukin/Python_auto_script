import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def generate_years(start_year: int, end_year: int):
    return [f"{y} – {y + 1}" for y in range(start_year, end_year)]


def set_cell_format(cell, text, align_center=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def process_docx(file_path: str, years_list: list):
    try:
        doc = Document(file_path)
        keywords = ["учебн", "год", "согласов", "рпд", "лист"]

        for table in doc.tables:
            if not table.rows:
                continue

            check_limit = min(len(table.rows), 3)
            rows_data = []
            for i in range(check_limit):
                text = " ".join(cell.text for cell in table.rows[i].cells).lower()
                rows_data.append(re.sub(r'\s+', ' ', text))

            full_header_text = " ".join(rows_data)
            matches = [kw for kw in keywords if kw in full_header_text]

            if len(matches) >= 2:
                target_row_index = 0
                for idx, text in enumerate(rows_data):
                    if "учебн" in text and "год" in text:
                        target_row_index = idx
                    elif "согласов" in text and target_row_index == 0:
                        target_row_index = idx

                for i in range(len(table.rows) - 1, target_row_index, -1):
                    row_el = table.rows[i]._element
                    row_el.getparent().remove(row_el)

                for year_str in years_list:
                    new_row = table.add_row()
                    set_cell_format(new_row.cells[0], year_str, align_center=True)
                    if len(new_row.cells) > 1:
                        set_cell_format(new_row.cells[1], "\n\n", align_center=False)

                doc.save(file_path)
                return True, "Успешно"

        return False, "Таблица не найдена"
    except Exception as e:
        return False, str(e)