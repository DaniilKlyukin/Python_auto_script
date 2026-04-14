import re
from docx import Document


def create_name_regex(name_str):
    parts = re.findall(r'[А-ЯЁа-яёA-Za-z]+', name_str)
    if not parts:
        return re.compile(re.escape(name_str), re.IGNORECASE)

    surname = parts[0] if len(parts[0]) > 2 else parts[-1]
    initials = [p[0] for p in parts if p != surname]

    if not initials:
        return re.compile(re.escape(surname), re.IGNORECASE)

    dots_spaces = r"\.?\s*"
    init_pattern = dots_spaces.join(initials) + r"\.?"

    pattern1 = rf"{init_pattern}\s*{surname}"
    pattern2 = rf"{surname}\s*{init_pattern}"

    return re.compile(f"({pattern1}|{pattern2})", re.IGNORECASE)


def create_title_regex(title_str):
    """Создает гибкое регулярное выражение для должности (учитывает сокращения)"""
    if not title_str:
        return None
    # Экранируем спецсимволы
    pattern = re.escape(title_str)
    # Позволяем слову 'заведующий' сокращаться до 'зав.'
    pattern = pattern.replace(r'заведующий', r'зав(едующий)?\.?')
    pattern = pattern.replace(r'Заведующий', r'[Зз]ав(едующий)?\.?')
    # Заменяем пробелы на возможность нескольких пробелов/переносов
    pattern = pattern.replace(r'\ ', r'\s+')
    return re.compile(pattern, re.IGNORECASE)


def is_signature_zone(paragraph, cell_context=False):
    text = paragraph.text.strip()
    if not text:
        return False

    if len(text) > 200:
        return False

    indicators = ["_", "20", "г.", "____________"]
    has_indicator = any(ind in text for ind in indicators)

    keywords = ["зав", "кафедр", "руковод", "программ", "декан", "председател", "составител", "разработчик"]

    para_content = text.lower()
    has_keyword = any(kw in para_content for kw in keywords)

    if cell_context:
        return True

    return has_indicator or has_keyword


def process_docx_signatures(file_path, old_name, new_name, old_title=None, new_title=None):
    try:
        doc = Document(file_path)
        name_regex = create_name_regex(old_name)
        title_regex = create_title_regex(old_title) if old_title else None
        is_changed = False

        def replace_in_paragraph(paragraph, is_cell=False):
            nonlocal is_changed
            # Проверяем, есть ли в параграфе ФИО или должность
            has_name = name_regex.search(paragraph.text)
            has_title = title_regex.search(paragraph.text) if title_regex else False

            if has_name or has_title:
                if is_signature_zone(paragraph, is_cell):
                    full_text = paragraph.text

                    # Сначала меняем должность, потом ФИО
                    if has_title and new_title:
                        full_text = title_regex.sub(new_title, full_text)

                    if has_name:
                        full_text = name_regex.sub(new_name, full_text)

                    if full_text != paragraph.text:
                        # Очищаем раны и записываем обновленный текст в первый ран
                        for i in range(len(paragraph.runs)):
                            paragraph.runs[i].text = ""
                        if paragraph.runs:
                            paragraph.runs[0].text = full_text
                        else:
                            paragraph.add_run(full_text)
                        is_changed = True

        for para in doc.paragraphs:
            replace_in_paragraph(para)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para, is_cell=True)

        if is_changed:
            doc.save(file_path)
            return True, "Успешно обновлено"
        else:
            return False, "Данные не найдены в подходящих блоках"

    except Exception as e:
        return False, str(e)