import os
import re
from docx import Document


def process_docx_signatures(file_path, old_fio, new_fio, old_title, new_title):
    """Замена ФИО и должности в тексте"""
    if not old_fio and not old_title:
        return False, "Нечего заменять"
    try:
        doc = Document(file_path)
        replaced = False
        for para in doc.paragraphs:
            if old_fio and old_fio in para.text:
                for run in para.runs:
                    if old_fio in run.text:
                        run.text = run.text.replace(old_fio, new_fio)
                        replaced = True
            if old_title and old_title in para.text:
                for run in para.runs:
                    if old_title in run.text:
                        run.text = run.text.replace(old_title, new_title)
                        replaced = True
        if replaced:
            doc.save(file_path)
            return True, "OK"
        return False, "не найдены"
    except Exception as e:
        return False, str(e)


def replace_word_in_docx(file_path, old_word, new_word):
    """Замена целого слова с сохранением форматирования"""
    try:
        doc = Document(file_path)
        pattern = re.compile(rf'\b{re.escape(old_word)}\b', re.IGNORECASE)
        replaced = 0

        all_paras = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paras.extend(cell.paragraphs)
        for section in doc.sections:
            all_paras.extend(section.header.paragraphs)
            all_paras.extend(section.footer.paragraphs)

        for para in all_paras:
            if not pattern.search(para.text):
                continue
            fmt = para.runs[0] if para.runs else None
            new_text = pattern.sub(new_word, para.text)
            if new_text != para.text:
                para.text = new_text
                if fmt:
                    for run in para.runs:
                        run.bold = fmt.bold
                        run.italic = fmt.italic
                        run.underline = fmt.underline
                replaced += 1

        if replaced:
            doc.save(file_path)
            return True, f"Замен: {replaced}"
        return False, "не найдено"
    except Exception as e:
        return False, str(e)


def main():
    print("=== Замена в .docx (рекурсивно) ===")
    dir_path = input("Путь к папке: ").strip().strip('"')
    if not os.path.isdir(dir_path):
        print("Ошибка: папка не найдена")
        return

    old_fio = input("Старое ФИО (Enter - пропустить): ").strip()
    new_fio = input("Новое ФИО: ").strip() if old_fio else ""

    old_title = input("Старая должность (Enter - пропустить): ").strip()
    new_title = input("Новая должность: ").strip() if old_title else ""

    old_word = input("Слово для замены (Enter - пропустить): ").strip()
    new_word = input("Новое слово: ").strip() if old_word else ""

    if not (old_fio or old_title or old_word):
        print("Ошибка: ничего не указано для замены")
        return

    stats = {"ok": 0, "skip": 0, "err": 0}

    for root, _, files in os.walk(dir_path):
        for f in files:
            if not f.lower().endswith('.docx') or f.startswith('~$'):
                continue

            path = os.path.join(root, f)
            rel = os.path.relpath(path, dir_path)
            modified = False

            if old_fio or old_title:
                ok, msg = process_docx_signatures(path, old_fio, new_fio, old_title, new_title)
                if ok:
                    modified = True
                elif "не найдены" not in msg:
                    print(f"[ERR] {rel}: {msg}")
                    stats["err"] += 1
                    continue

            if old_word:
                ok, msg = replace_word_in_docx(path, old_word, new_word)
                if ok:
                    modified = True
                elif "не найдено" not in msg:
                    print(f"[ERR] {rel}: {msg}")
                    stats["err"] += 1
                    continue

            if modified:
                print(f"[OK] {rel}")
                stats["ok"] += 1
            else:
                print(f"[SKIP] {rel}")
                stats["skip"] += 1

    print(f"\n=== Готово ===")
    print(f"Обработано: {stats['ok']}, Пропущено: {stats['skip']}, Ошибок: {stats['err']}")


if __name__ == "__main__":
    main()