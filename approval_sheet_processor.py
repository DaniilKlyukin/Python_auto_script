import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def generate_years(start_year, end_year):
    years_list = []
    for y in range(start_year, end_year):
        years_list.append(f"{y} – {y + 1}")
    return years_list


def set_cell_format(cell, text, align_center=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)


def process_docx(file_path, years_list):
    try:
        # Ошибка возникает именно здесь, если файл "битый" внутри
        doc = Document(file_path)
        table_updated = False

        for table in doc.tables:
            if len(table.rows) == 0:
                continue

            # Улучшенный поиск: проверяем все ячейки первой строки
            # (иногда "учебный год" может быть во 2-м столбце или ячейки объединены)
            first_row_text = " ".join([cell.text.lower() for cell in table.rows[0].cells])

            if "учебный год" in first_row_text:
                # Удаляем старые строки (кроме шапки)
                # Идем с конца, чтобы индексы не смещались
                rows = table.rows
                for i in range(len(rows) - 1, 0, -1):
                    row_el = rows[i]._element
                    row_el.getparent().remove(row_el)

                # Добавляем новые
                for year_str in years_list:
                    new_row = table.add_row()
                    set_cell_format(new_row.cells[0], year_str, align_center=True)
                    if len(new_row.cells) > 1:
                        set_cell_format(new_row.cells[1], "\n\n", align_center=False)

                table_updated = True
                break

        if table_updated:
            doc.save(file_path)
            return True, "Успешно"
        else:
            return False, "Таблица 'Лист согласования' не найдена"

    except KeyError as e:
        return False, f"Ошибка структуры (отсутствует медиа-файл): {e}"
    except Exception as e:
        return False, f"Ошибка: {e}"


def main():
    print("=== Обновление лет в Листах согласования ===")
    folder_path = input("Введите путь до папки: ").strip('\'"')

    if not os.path.isdir(folder_path):
        print("Ошибка: Путь не найден.")
        return

    try:
        start_year = int(input("Введите начальный год: "))
        end_year = int(input("Введите конечный год: "))
    except ValueError:
        return

    years_list = generate_years(start_year, end_year)
    processed_count = 0
    error_count = 0

    for filename in os.listdir(folder_path):
        if filename.lower().endswith('.docx') and not filename.startswith('~$'):
            file_path = os.path.join(folder_path, filename)
            print(f"Обработка: {filename}...", end=" ")

            success, message = process_docx(file_path, years_list)

            if success:
                print("ОК!")
                processed_count += 1
            else:
                print(f"ПРОПУЩЕНО ({message})")
                error_count += 1

    print(f"\nИтог: Обновлено: {processed_count}, Пропущено: {error_count}")


if __name__ == "__main__":
    main()