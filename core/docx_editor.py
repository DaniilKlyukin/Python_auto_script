import os
import tempfile
import fitz
from docx import Document
from docx.shared import Mm
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


class DocxEditor:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.doc = None
        self._temp_images = []

    def __enter__(self):
        self.doc = Document(self.file_path)
        return self

    def __exit__(self, exc_type, exc_val, exc_tb):
        if self.doc and exc_type is None:
            self.doc.save(self.file_path)
        for tmp in self._temp_images:
            try: os.remove(tmp)
            except: pass

    def _prepare_image(self, path: str) -> str:
        if path.lower().endswith('.pdf'):
            doc = fitz.open(path)
            pix = doc.load_page(0).get_pixmap(dpi=300)
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
            pix.save(tmp.name)
            doc.close()
            self._temp_images.append(tmp.name)
            return tmp.name
        return path

    def _make_floating(self, shape, w_mm, h_mm):
        w_emu, h_emu = int(w_mm * 36000), int(h_mm * 36000)
        inline = shape._inline
        anchor_xml = f"""
        <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="251658240" 
            behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1" {nsdecls('wp', 'wp14', 'pic', 'r')}>
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="page"><wp:posOffset>0</wp:posOffset></wp:positionV>
            <wp:extent cx="{w_emu}" cy="{h_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/><wp:docPr id="1" name="Scan"/><wp:cNvGraphicFramePr/>
            {inline.graphic.xml}
        </wp:anchor>
        """
        inline.getparent().replace(inline, parse_xml(anchor_xml))

    def add_floating_scan(self, paragraph, image_path, width_mm=210, height_mm=297):
        img_src = self._prepare_image(image_path)
        run = paragraph.add_run()
        shape = run.add_picture(img_src, width=Mm(width_mm))
        self._make_floating(shape, width_mm, height_mm)

    def add_image_at_beginning(self, image_path: str):
        self.add_floating_scan(self.doc.paragraphs[0], image_path)

    def add_image_at_end(self, image_path: str):
        self.add_floating_scan(self.doc.add_paragraph(), image_path)

    def insert_image_after_text(self, text: str, image_path: str) -> bool:
        for p in self.doc.paragraphs:
            if text.lower() in p.text.lower():
                self.add_floating_scan(p, image_path)
                return True
        for t in self.doc.tables:
            for r in t.rows:
                for c in r.cells:
                    if text.lower() in c.text.lower():
                        self.add_floating_scan(c.paragraphs[0], image_path)
                        return True
        return False